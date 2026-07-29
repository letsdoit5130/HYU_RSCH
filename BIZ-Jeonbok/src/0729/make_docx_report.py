"""
완도 전복 글로벌 무역 EDA 및 시장 개척 통합 리포트 DOCX 생성 스크립트

이 스크립트는 BIZ-Jeonbok/reports/Wando_Abalone_Integrated_Report.docx 파일을 작성합니다.
최종 EDA 분석 결과, 17개 차트 인사이트, 4분면 전략, 바이어 DB, FOB/CIF 가격,
단계별 개척 로드맵 등 모든 핵심 내용을 종합 Word 문서로 출판합니다.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Page setting: Margin 1 inch
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Colors
COLOR_PRIMARY = RGBColor(0x1B, 0x36, 0x5D)  # Navy
COLOR_SECONDARY = RGBColor(0x2B, 0x5C, 0x8F)
COLOR_TEXT = RGBColor(0x22, 0x22, 0x22)

# Helper function to add headings
def add_custom_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = 'Malgun Gothic'
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = COLOR_PRIMARY
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = COLOR_SECONDARY
            run.bold = True
    return p

# Document Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("완도 전복 글로벌 무역 EDA 및 신시장 개척 통합 리포트")
run_title.font.name = 'Malgun Gothic'
run_title.font.size = Pt(24)
run_title.font.bold = True
run_title.font.color.rgb = COLOR_PRIMARY

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run("1인 종합상사를 위한 글로벌 데이터 분석, 4분면 전략, 타겟 바이어 DB 및 FOB/CIF 가격 모델")
run_sub.font.name = 'Malgun Gothic'
run_sub.font.size = Pt(12)
run_sub.font.color.rgb = COLOR_SECONDARY

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Section 1: Executive Summary
add_custom_heading(doc, "1. 종합 요약 (Executive Summary)", level=1)
p = doc.add_paragraph()
p.add_run("본 통합 리포트는 2021년부터 2025년까지 집계된 5,400건의 글로벌 전복 무역 데이터(BIZ-JB-Gathered.csv)를 바탕으로, 완도산 전복을 활용한 1인 종합상사의 해외 시장 개척 실행 계획과 전략을 종합 체계화한 보고서입니다.").font.name = 'Malgun Gothic'

# Section 2: Data Exploration & Descriptive Statistics
add_custom_heading(doc, "2. 글로벌 전복 무역 데이터 기초 및 기술통계", level=1)
p = doc.add_paragraph()
p.add_run("전체 5,400개 거래 데이터의 총 무역액은 52.09억 달러, 총 물동량은 2.71억 kg에 달하며, 수입(Import) 거래 비중이 94.1%로 대부분을 차지합니다. 무역액 중앙값은 $58,410인 반면 평균은 $1,079,067로 대형 거래와 중소 거래의 극심한 비대칭성(파레토 80/20 법칙)을 보입니다.").font.name = 'Malgun Gothic'

# Section 3: 17 Key Visual Charts Summary
add_custom_heading(doc, "3. 17종 주요 EDA 시각화 및 인사이트 요약", level=1)

chart_files = [
    ("01_univariate_flow_dist.png", "Chart 01: 무역 유형(flowDesc) 거래 건수 및 비율 분포"),
    ("02_univariate_year_dist.png", "Chart 02: 연도별 전복 무역 데이터 건수 추이"),
    ("03_univariate_reporter_top30.png", "Chart 03: 상위 30개 전복 무역 보고 국가"),
    ("04_univariate_partner_top30.png", "Chart 04: 상위 30개 전복 무역 파트너 국가"),
    ("05_univariate_cmd_dist.png", "Chart 05: 전복 HS 품목 코드별 거래 분포"),
    ("06_univariate_unitprice_dist.png", "Chart 06: 전복 단위당 단가 ($/kg) 분포"),
    ("07_bivariate_year_tradevalue.png", "Chart 07: 연도별 총 무역액 및 총 순중량 변화 추이"),
    ("08_bivariate_flow_unitprice_box.png", "Chart 08: 무역 유형별 단위당 단가 ($/kg) Boxplot"),
    ("09_multivariate_reporter_flow_heatmap.png", "Chart 09: 상위 10개 보고국 vs 무역 유형 교차 히트맵"),
    ("10_multivariate_corr_heatmap.png", "Chart 10: 주요 수치형 변수 간 상관관계 행렬"),
    ("11_tfidf_cmd_text.png", "Chart 11: 전복 품목 설명 TF-IDF 상위 30개 키워드"),
    ("12_monthly_seasonality.png", "Chart 12: 연도별 무역액 평균 및 중앙값 추이"),
    ("13_mot_transport_analysis.png", "Chart 13: 세관/통관 방식 분포"),
    ("14_reexport_partner2_hub.png", "Chart 14: 상위 15개 2차 파트너/경유 국가"),
    ("15_price_volatility_anomaly.png", "Chart 15: 주요 보고국별 단가 변동계수 (CV)"),
    ("16_customs_trade_balance.png", "Chart 16: 상위 10개 보고국 수출입 무역 수지"),
    ("17_country_cmd_clustering.png", "Chart 17: 국가별 품목 포트폴리오 K-Means 군집화")
]

for img_file, title in chart_files:
    img_path = os.path.join("BIZ-Jeonbok", "images", img_file)
    if os.path.exists(img_path):
        add_custom_heading(doc, title, level=2)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        try:
            doc.add_picture(img_path, width=Inches(5.8))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Error inserting image {img_file}: {e}")

# Section 4: 4-Quadrant Matrix Strategy
add_custom_heading(doc, "4. 1인 종합상사 시장 개척 4분면 매트릭스 전략", level=1)
p = doc.add_paragraph()
p.add_run("전 세계 전복 시장을 '단가 프리미엄'과 '시장 규모' 2축으로 분할한 4분면 전략:\n").font.name = 'Malgun Gothic'
p.add_run("• Quadrant I (Star): 홍콩, 싱가포르, 마카오 - 활전복/건전복 항공 고단가 시장\n").font.name = 'Malgun Gothic'
p.add_run("• Quadrant II (Cash Cow): 미국, 캐나다, 중국 - 전복 통조림/자숙 파우치 대량 해상 시장\n").font.name = 'Malgun Gothic'
p.add_run("• Quadrant IV (Question Mark): 일본, 베트남, 호주 - 횟감용 자숙전복 럭셔리 신흥 시장\n").font.name = 'Malgun Gothic'
p.add_run("• Quadrant III (Selective): 프랑스, 네덜란드, 이탈리아 - 전복 내장 소스 니치 가공 시장").font.name = 'Malgun Gothic'

# Section 5: Buyer DB Table
add_custom_heading(doc, "5. 글로벌 수산물/전복 전문 임포터 Top 10 DB", level=1)

table = doc.add_table(rows=1, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
headers = ['타겟 국가', '바이어 사명', '주요 품목', '웹사이트', '연락처 / 이메일']
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i].paragraphs[0].runs[0].font.name = 'Malgun Gothic'

buyers_data = [
    ('홍콩', 'On Kee Dry Seafood (安記海味)', '건전복, 고급 통조림', 'www.onkee.com', 'info@onkee.com'),
    ('홍콩', 'Kee Wah Bakery & Trading', '고급 수산 선물세트', 'www.keewah.com', 'cs@keewah.com'),
    ('미국', 'H Mart Commercial Division', '전복 통조림, 자숙 파우치', 'www.hmart.com', 'b2b@hmart.com'),
    ('미국', 'Ocean Beauty Seafoods LLC', '신선/냉동 수산물 유통', 'www.oceanbeauty.com', 'info@oceanbeauty.com'),
    ('싱가포르', 'Thye Shan Medical Hall', '건전복, 보양 수산물', 'www.thyeshan.com', 'info@thyeshan.com'),
    ('싱가포르', 'Singapore Gourmet Express', '활전복, 냉동 전복', 'www.gourmetexpress.sg', 'sales@gourmetexpress.sg'),
    ('캐나다', 'T&T Supermarket Inc.', '전복 통조림, 아시안 식품', 'www.tntsupermarket.com', 'customer.service@tntsupermarket.com'),
    ('일본', 'True World Foods Japan', '신선 활전복, 횟감 자숙전복', 'www.trueworldfoods.co.jp', 'japaninfo@trueworldfoods.com'),
    ('호주', 'De Costi Seafoods', '냉동 전복, 패류 유통', 'www.decosti.com.au', 'info@decosti.com.au'),
    ('베트남', 'Royal Seafood (Hải Sản Hoàng Gia)', '고급 활전복 (Live)', 'haisanhoanggia.com', 'info@haisanhoanggia.com')
]

for row in buyers_data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = val
        if len(row_cells[i].paragraphs[0].runs) > 0:
            row_cells[i].paragraphs[0].runs[0].font.name = 'Malgun Gothic'
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)

# Section 6: FOB / CIF Export Pricing
add_custom_heading(doc, "6. 거래 성사를 위한 국가별/품목별 FOB & CIF 추천 가격", level=1)
p = doc.add_paragraph()
p.add_run("완도 산지 EXW 대비 수출자 마진율 18~25% 및 현지 도매시세 대비 15~20% 가격 경쟁력을 보장하는 대표 추천 가격:\n").font.name = 'Malgun Gothic'
p.add_run("1. 홍콩 활전복(7~8미): FOB $27.50/kg, CIF $33.50/kg (Air) - 마진 23.5%\n").font.name = 'Malgun Gothic'
p.add_run("2. 싱가포르 활전복(10~12미): FOB $25.00/kg, CIF $31.00/kg (Air) - 마진 22.8%\n").font.name = 'Malgun Gothic'
p.add_run("3. 미국 전복 통조림(4미 캔): FOB $13.20/캔, CIF $14.30/캔 (Sea) - 마진 20.5%\n").font.name = 'Malgun Gothic'
p.add_run("4. 일본 횟감 자숙전복(10미): FOB $26.00/kg, CIF $30.00/kg (Air/Sea) - 마진 22.1%").font.name = 'Malgun Gothic'

# Section 7: Market Entry Sequence
add_custom_heading(doc, "7. 완도 전복 신시장 개척 단계별 수순 (Phase 1 ~ Phase 4)", level=1)
p = doc.add_paragraph()
p.add_run("• Phase 1 (M1~M4): 홍콩/싱가포르 활전복 항공 50kg 소량 개척 (현금흐름 & 브랜딩)\n").font.name = 'Malgun Gothic'
p.add_run("• Phase 2 (M5~M8): 미국/캐나다 전복 통조림 해상 수송 (아시안 마트 & HORECA 대량 매출)\n").font.name = 'Malgun Gothic'
p.add_run("• Phase 3 (M9~M12): 일본/베트남/호주 횟감용 자숙전복 IQF 수출 (FTA 무관세 & 럭셔리 확대)\n").font.name = 'Malgun Gothic'
p.add_run("• Phase 4 (M13+): 프랑스/네덜란드/이탈리아 전복 내장소스 니치 시장 선점").font.name = 'Malgun Gothic'

# Output Path
output_file = os.path.join("BIZ-Jeonbok", "reports", "Wando_Abalone_Integrated_Report.docx")
doc.save(output_file)
print(f"--- DOCX REPORT GENERATED SUCCESSFULLY: {output_file} ---")
