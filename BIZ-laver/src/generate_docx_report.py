"""
수출 EDA 종합 및 심화 전략 보고서 Word (.docx) v2 파일 생성 스크립트.
"""

import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

DOCX_OUTPUT_PATH = 'BIZ-laver/reports/HaeYu_Laver_Export_EDA_and_Market_Strategy_Report_v2.docx'
IMAGE_DIR = 'BIZ-laver/images'

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(26, 82, 118)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(46, 134, 193)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(40, 116, 166)
    return p

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(12)
t_run = title_p.add_run("해유 김 수출 EDA 종합 및 1인 무역회사 심화 시장개척 전략 보고서 (v2)")
t_run.font.name = '맑은 고딕'
t_run.font.size = Pt(22)
t_run.font.bold = True
t_run.font.color.rgb = RGBColor(26, 82, 118)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(24)
s_run = sub_p.add_run("2021~2025 글로벌 김 수출 데이터 분석, 단가 마진 구간 및 4분면 포트폴리오 분석")
s_run.font.name = '맑은 고딕'
s_run.font.size = Pt(12)
s_run.font.color.rgb = RGBColor(120, 144, 156)

add_styled_heading(doc, "1. Executive Summary", 1)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("본 보고서는 2021년부터 2025년까지 대한민국 김 수출 데이터(1,069건)를 바탕으로 Exploratory Data Analysis(EDA)를 수행하고, 1인 무역회사 관점에서 최적의 신규 유망 시장 개척 전략 및 마진 포트폴리오를 도출하였습니다.")
r.font.name = '맑은 고딕'
r.font.size = Pt(10.5)

add_styled_heading(doc, "2. 추가 심화 EDA 분석 및 마진/포트폴리오 평가", 1)

img12_path = os.path.join(IMAGE_DIR, '12_advanced_monthly_seasonality.png')
if os.path.exists(img12_path):
    doc.add_paragraph("그림 1: 연도별 품목별 총 수출액 추이 (2021~2025)")
    doc.add_picture(img12_path, width=Inches(6.0))

p = doc.add_paragraph()
r = p.add_run("조미김(HS 200899)은 2025년 평균 수출 단가가 $30.41/kg에 도달하며 마른김($21.56/kg) 대비 +41.1%의 고마진 프리미엄을 형성하였습니다.")
r.font.name = '맑은 고딕'

table1 = doc.add_table(rows=6, cols=7)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["연도", "마른김 수출액", "마른김 물량(t)", "마른김 단가", "조미김 수출액", "조미김 물량(t)", "조미김 단가"]
for i, h in enumerate(headers):
    cell = table1.cell(0, i)
    cell.text = h
    set_cell_background(cell, "1A5274")
    p = cell.paragraphs[0]
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.bold = True

data1 = [
    ["2021년", "$253.74M", "24,419.5 t", "$17.88/kg", "$553.99M", "25,297.5 t", "$25.45/kg"],
    ["2022년", "$273.17M", "25,300.8 t", "$20.19/kg", "$476.95M", "24,365.8 t", "$21.64/kg"],
    ["2023년", "$354.88M", "28,158.4 t", "$18.42/kg", "$547.15M", "26,734.2 t", "$22.35/kg"],
    ["2024년", "$476.56M", "29,680.8 t", "$20.34/kg", "$636.48M", "26,955.3 t", "$27.11/kg"],
    ["2025년", "$563.58M", "31,147.0 t", "$21.56/kg", "$687.40M", "27,502.0 t", "$30.41/kg"],
]

for row_idx, row_data in enumerate(data1, start=1):
    for col_idx, text in enumerate(row_data):
        cell = table1.cell(row_idx, col_idx)
        cell.text = text
        if row_idx % 2 == 1:
            set_cell_background(cell, "F4F6F7")

img13_path = os.path.join(IMAGE_DIR, '13_advanced_price_bracket_distribution.png')
if os.path.exists(img13_path):
    doc.add_paragraph("\n그림 2: 품목별 수출 단가($/kg) 5대 마진 구간 비중 (%)")
    doc.add_picture(img13_path, width=Inches(6.0))

img14_path = os.path.join(IMAGE_DIR, '14_advanced_market_portfolio_matrix.png')
if os.path.exists(img14_path):
    doc.add_paragraph("\n그림 3: 글로벌 40개국 시장 포트폴리오 4분면 매트릭스")
    doc.add_picture(img14_path, width=Inches(6.0))

img15_path = os.path.join(IMAGE_DIR, '15_advanced_seasoned_vs_raw_monthly_trend.png')
if os.path.exists(img15_path):
    doc.add_paragraph("\n그림 4: 주요 수출 대상국 단가 변동성 (CV % 리스크)")
    doc.add_picture(img15_path, width=Inches(6.0))

add_styled_heading(doc, "3. 1인 무역회사 최종 시장개척 액션 플랜", 1)
p = doc.add_paragraph()
r = p.add_run("1인 무역회사는 조미김(HS 200899) 스낵 제품군에 100% 집중하여, Star Market인 UAE, 폴란드, 콜롬비아, 튀르키예와 Rising Volume 시장인 사우디아라비아, 카자흐스탄을 주력 개척 국가로 선정해야 합니다.")
r.font.name = '맑은 고딕'
r.font.bold = True

doc.save(DOCX_OUTPUT_PATH)
print(f"Word 보고서가 {DOCX_OUTPUT_PATH}에 성공적으로 업데이트 저장되었습니다.")
