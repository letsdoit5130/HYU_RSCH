"""
1인 종합상사 전복 수출 시장 개척 실전 가이드 Word(.docx) 보고서 생성 스크립트

이 스크립트는 BIZ-Jeonbok/artifacts/Solo_Trader_Abalone_Export_Guide.md 리포트 내용과
생성된 시각화 차트 이미지들을 기반으로 python-docx를 사용하여
제목, 목차, 스타일링 표, 서식, 차트 이미지가 정밀 조합된 전문 Word 문서(.docx)를 생성합니다.
"""
import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ----------------------------------------------------
# 0. 경로 설정
# ----------------------------------------------------
BASE_DIR = 'BIZ-Jeonbok'
MD_PATH = os.path.join(BASE_DIR, 'artifacts', 'Solo_Trader_Abalone_Export_Guide.md')
IMAGES_DIR = os.path.join(BASE_DIR, 'images')
REPORTS_DOCX = os.path.join(BASE_DIR, 'reports', 'Solo_Trader_Abalone_Export_Guide.docx')
ARTIFACTS_DOCX = os.path.join(BASE_DIR, 'artifacts', 'Solo_Trader_Abalone_Export_Guide.docx')

# ----------------------------------------------------
# 1. Word Document 및 스타일 설정
# ----------------------------------------------------
doc = Document()

sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style_normal = doc.styles['Normal']
font = style_normal.font
font.name = '맑은 고딕'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    run = h.runs[0]
    run.font.name = '맑은 고딕'
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x66)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return h

# ----------------------------------------------------
# 2. 문서 타이틀 작성
# ----------------------------------------------------
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(20)
title_p.paragraph_format.space_after = Pt(10)
title_run = title_p.add_run("1인 종합상사를 위한 전복 수출 종합 전략 & HS CODE별 유망국가 TOP 10 가이드")
title_run.font.name = '맑은 고딕'
title_run.font.size = Pt(18)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(20)
sub_run = sub_p.add_run("― 최적 아이템x시장 콤보 전략 및 HS CODE별 유망국가 TOP 10 실적/단가 분석 ―")
sub_run.font.name = '맑은 고딕'
sub_run.font.size = Pt(11)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# ----------------------------------------------------
# 3. Markdown 파싱
# ----------------------------------------------------
with open(MD_PATH, 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_table = False
table_lines = []

def parse_markdown_table(lines):
    rows = []
    for line in lines:
        if '|' in line and not line.strip().startswith('|---') and not line.strip().startswith('|:--'):
            parts = [p.strip() for p in line.strip().split('|')[1:-1]]
            if any(parts):
                rows.append(parts)
    return rows

for line in lines:
    line_str = line.strip()
    
    if line_str == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        continue

    if line_str.startswith('# '):
        continue
    elif line_str.startswith('## '):
        add_styled_heading(doc, line_str[3:], 1)
        continue
    elif line_str.startswith('### '):
        add_styled_heading(doc, line_str[4:], 2)
        continue
    elif line_str.startswith('#### '):
        add_styled_heading(doc, line_str[5:], 3)
        continue

    img_match = re.search(r'!\[.*?\]\(\.\./images/(.*?)\)', line_str)
    if img_match:
        img_name = img_match.group(1)
        img_full_path = os.path.join(IMAGES_DIR, img_name)
        if os.path.exists(img_full_path):
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.paragraph_format.space_before = Pt(8)
            img_p.paragraph_format.space_after = Pt(8)
            run = img_p.add_run()
            run.add_picture(img_full_path, width=Inches(5.8))
            
            caption_p = doc.add_paragraph()
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_p.paragraph_format.space_after = Pt(10)
            c_run = caption_p.add_run(f"[시각화 차트: {img_name}]")
            c_run.font.size = Pt(9)
            c_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        continue

    if '|' in line_str:
        in_table = True
        table_lines.append(line_str)
        continue
    else:
        if in_table:
            parsed_rows = parse_markdown_table(table_lines)
            if parsed_rows:
                table = doc.add_table(rows=len(parsed_rows), cols=len(parsed_rows[0]))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True
                
                for r_idx, row_data in enumerate(parsed_rows):
                    row = table.rows[r_idx]
                    for c_idx, cell_value in enumerate(row_data):
                        cell = row.cells[c_idx]
                        cell.text = cell_value
                        
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(3)
                        p.paragraph_format.space_after = Pt(3)
                        for r in p.runs:
                            r.font.name = '맑은 고딕'
                            r.font.size = Pt(9)
                        
                        if r_idx == 0:
                            set_cell_background(cell, '1B365D')
                            for r in p.runs:
                                r.font.bold = True
                                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        else:
                            if r_idx % 2 == 1:
                                set_cell_background(cell, 'F2F4F7')
            
            doc.add_paragraph()
            table_lines = []
            in_table = False

    if line_str.startswith('> '):
        quote_p = doc.add_paragraph()
        quote_p.paragraph_format.left_indent = Inches(0.4)
        quote_p.paragraph_format.space_before = Pt(3)
        quote_p.paragraph_format.space_after = Pt(3)
        q_run = quote_p.add_run(line_str[2:])
        q_run.font.name = '맑은 고딕'
        q_run.font.size = Pt(9.5)
        q_run.font.bold = True
        q_run.font.color.rgb = RGBColor(0x00, 0x4D, 0x40)
        continue

    if line_str.startswith('- '):
        list_p = doc.add_paragraph(style='List Bullet')
        list_p.paragraph_format.space_before = Pt(2)
        list_p.paragraph_format.space_after = Pt(2)
        
        text_content = line_str[2:]
        parts = re.split(r'(\*\*.*?\*\*)', text_content)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = list_p.add_run(part[2:-2])
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            else:
                list_p.add_run(part)
        continue

    if line_str:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.25
        
        parts = re.split(r'(\*\*.*?\*\*)', line_str)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2])
                r.font.bold = True
            else:
                p.add_run(part)

# ----------------------------------------------------
# 4. 안전 저장
# ----------------------------------------------------
try:
    doc.save(ARTIFACTS_DOCX)
    print(f"Artifacts Word 저장 성공: {ARTIFACTS_DOCX}")
except Exception as e:
    print(f"Artifacts Word 저장 경고: {e}")

try:
    doc.save(REPORTS_DOCX)
    print(f"Reports Word 저장 성공: {REPORTS_DOCX}")
except Exception as e:
    alt_reports_path = os.path.join(BASE_DIR, 'reports', 'Solo_Trader_Abalone_Export_Guide_v3.docx')
    try:
        doc.save(alt_reports_path)
        print(f"Reports Word v3 저장 성공: {alt_reports_path}")
    except Exception as e2:
        print(f"Word 저장 건너뜀 (이미 락 또는 다른 프로세스 오픈 상태): {e2}")
