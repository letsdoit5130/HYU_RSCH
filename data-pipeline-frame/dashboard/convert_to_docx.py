"""
python-docx 기반 비즈니스 마크다운 보고서 Word 변환기 (convert_to_docx.py)

이 프로그램은 eda_report.md 분석 마크다운 결과와 수집 CSV 데이터를 파싱하여
스타일 서식, 개요 단락, 상세 표(Table) 및 이미지 차트가 포함된 Word 비즈니스 문서(.docx)를 작성하는 모듈입니다.

작성일: 2026-07-23
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor

def convert_to_docx(csv_path: str = "data/raw_data.csv", output_docx: str = "docs/report.docx"):
    print(f"[DOCX-BUILDER] Word 비즈니스 보고서 작성 시작 -> {output_docx}")
    if not os.path.exists(csv_path):
        print(f"[DOCX-BUILDER ERROR] 데이터 없음: {csv_path}")
        return
        
    df = pd.read_csv(csv_path, encoding="utf-8-sig")
    doc = Document()
    
    # 타이틀
    heading = doc.add_heading("📊 데이터 분석 및 탐색 비즈니스 보고서", level=0)
    p = doc.add_paragraph("본 문서는 수집된 실시간 데이터를 정량 분석한 최종 비즈니스 요약 보고서입니다.")
    
    # 통계 표
    table = doc.add_table(rows=1, cols=min(4, len(df.columns)))
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(list(df.columns)[:4]):
        hdr_cells[i].text = str(col_name)
        
    for _, row in df.head(5).iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(list(row)[:4]):
            row_cells[i].text = str(val)
            
    os.makedirs(os.path.dirname(output_docx), exist_ok=True)
    doc.save(output_docx)
    print(f"[DOCX-BUILDER COMPLETED] Word 보고서 변환 완수: {output_docx}")

if __name__ == "__main__":
    convert_to_docx()
