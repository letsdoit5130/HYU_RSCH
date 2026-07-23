"""
이 스크립트는 yes24/docs/eda_report.md 마크다운 파일을 로드하여
동일 디렉토리에 동일 명칭의 고품질 Word 문서(eda_report.docx) 파일로 변환하여 생성하는 프로그램입니다.
주요 기능:
- 마크다운 헤더(#, ##, ###)를 Word 문서 스타일 제목(Heading 1~3)으로 변환
- 마크다운 내 이미지 링크(![라벨](경로))를 감지하여 실제 이미지 파일을 문서 내에 적절한 크기로 임베딩
- 마크다운 데이터 표(| 컬럼 |) 구조를 파싱하여 배경색(헤더)과 격자 테두리가 있는 워드 표 객체로 재구성
- 본문 텍스트, 인용구( 들여쓰기 및 이탤릭 적용), 코드 블록 등의 상세 서식 구현
"""
# -*- coding: utf-8 -*-
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# 1. 경로 설정
base_dir = r"C:\Users\leeak\OneDrive\1.HaeYu\HYU_RSCH\yes24"
md_path = os.path.join(base_dir, "docs", "eda_report.md")
docx_path = os.path.join(base_dir, "docs", "eda_report.docx")
image_base_dir = os.path.join(base_dir, "images")

# 2. 스타일 헬퍼 함수 정의
def set_cell_background(cell, hex_color):
    """셀의 배경색을 지정된 16진수 색상 코드로 설정합니다."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """셀 내부 여백(패딩)을 DXA 단위로 설정합니다."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table):
    """테이블 전체 테두리를 연한 회색 실선으로 설정합니다."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def parse_markdown_table_row(line):
    """마크다운 테이블 행 텍스트를 파싱하여 셀 목록을 반환합니다."""
    # 앞뒤 파이프(|) 문자 제거 후 스플릿
    stripped = line.strip().strip('|')
    if not stripped:
        return []
    parts = stripped.split('|')
    return [p.strip() for p in parts]

def convert_md_to_docx():
    # 마크다운 파일 열기
    if not os.path.exists(md_path):
        print(f"오류: 마크다운 파일이 존재하지 않습니다: {md_path}")
        return

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # 페이지 크기 및 여백 기본 설정 (A4 규격 적용)
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 가로 너비
    section.page_height = Inches(11.69) # A4 세로 높이
    section.top_margin = Inches(1.0)   # 사방 1인치 여백
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # 기본 글꼴 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10)
    font.color.rgb = RGBColor(51, 51, 51) # 진한 회색

    # 문단 행간 조절
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after = Pt(6)

    # 상태 관리 변수
    in_table = False
    table_data = []
    in_code_block = False
    code_text = ""

    i = 0
    while i < len(lines):
        line = lines[i]
        line_stripped = line.strip()

        # 1. 코드 블록 처리
        if line_stripped.startswith("```"):
            if in_code_block:
                # 코드 블록 종료
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text.strip())
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(100, 100, 100)
                # 배경 연한 회색 음영
                pPr = p._p.get_or_add_pPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                pPr.append(shd)

                in_code_block = False
                code_text = ""
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_text += line
            i += 1
            continue

        # 2. 마크다운 테이블(표) 감지 및 파싱
        if line_stripped.startswith("|"):
            # 헤더 구분선인 경우 (예: |:---|---|) 다음 행으로 스킵
            if re.match(r'^\|[\s:-|]*\|$', line_stripped):
                i += 1
                continue

            in_table = True
            row_cells = parse_markdown_table_row(line)
            if row_cells:
                table_data.append(row_cells)
            i += 1
            continue
        else:
            # 테이블 수집이 끝나고 일반 문단을 만나면 누적된 테이블 빌드
            if in_table:
                if table_data:
                    # 열 개수 감지
                    col_count = max(len(row) for row in table_data)
                    # Word 표 추가
                    table = doc.add_table(rows=0, cols=col_count)
                    set_table_borders(table)

                    for r_idx, row_data in enumerate(table_data):
                        # 셀 데이터 보강 (열 개수 맞춤)
                        while len(row_data) < col_count:
                            row_data.append("")

                        row_cells = table.add_row().cells
                        for c_idx, val in enumerate(row_data):
                            cell = row_cells[c_idx]
                            cell.text = val
                            
                            # 내부 여백 설정
                            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

                            # 글자 포맷 설정 (헤더인 경우 굵게 및 배경색 부여)
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(2)
                            run = p.runs[0] if p.runs else p.add_run()
                            run.font.name = 'Malgun Gothic'
                            run.font.size = Pt(9.5)

                            if r_idx == 0:
                                # 헤더행 디자인
                                run.font.bold = True
                                set_cell_background(cell, "D5E8F0") # 연한 푸른빛 헤더
                            elif r_idx % 2 == 1:
                                # 홀수행 번갈아가며 연한 회색 음영 배경 처리 (가독성 증대)
                                set_cell_background(cell, "FAFAFA")

                    # 테이블 아래 한 줄 간격 확보
                    doc.add_paragraph().paragraph_format.space_before = Pt(6)

                in_table = False
                table_data = []

        # 3. 빈 라인 처리
        if not line_stripped:
            i += 1
            continue

        # 4. 이미지 링크 처리 (![라벨](경로))
        img_match = re.search(r'!\[(.*?)\]\((.*?)\)', line_stripped)
        if img_match:
            img_label = img_match.group(1)
            img_rel_path = img_match.group(2)
            
            # 경로 유연하게 변환 (eda_report.md에 "images/xxx.png"로 기록됨)
            filename = os.path.basename(img_rel_path)
            img_file_path = os.path.join(image_base_dir, filename)

            if os.path.exists(img_file_path):
                # 이미지 삽입 (너비 약 5.8인치 설정)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_file_path, width=Inches(5.8))
                
                # 이미지 설명 캡션 추가
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_before = Pt(4)
                caption.paragraph_format.space_after = Pt(12)
                cap_run = caption.add_run(f"▲ 그림: {img_label}")
                cap_run.font.size = Pt(8.5)
                cap_run.font.italic = True
                cap_run.font.color.rgb = RGBColor(120, 120, 120)
            else:
                # 이미지가 없는 경우 단순 텍스트 표시
                p = doc.add_paragraph()
                p.add_run(f"[그림 대체 텍스트: {img_label} ({filename} 파일을 찾을 수 없습니다)]").font.color.rgb = RGBColor(200, 100, 100)
            i += 1
            continue

        # 5. 헤더(제목) 스타일 처리 (#, ##, ###)
        if line_stripped.startswith("###"):
            title = line_stripped.replace("###", "").strip()
            p = doc.add_heading(title, level=3)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.runs[0]
            run.font.name = 'Malgun Gothic'
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 30, 30) # 어두운 회색
            i += 1
            continue
        elif line_stripped.startswith("##"):
            title = line_stripped.replace("##", "").strip()
            p = doc.add_heading(title, level=2)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            run = p.runs[0]
            run.font.name = 'Malgun Gothic'
            run.font.bold = True
            run.font.color.rgb = RGBColor(46, 117, 182) # 청화색 테마 적용
            i += 1
            continue
        elif line_stripped.startswith("#"):
            title = line_stripped.replace("#", "").strip()
            p = doc.add_heading(title, level=1)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            run = p.runs[0]
            run.font.name = 'Malgun Gothic'
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0) # 검은색
            i += 1
            continue

        # 6. 인용구 처리 (>)
        if line_stripped.startswith(">"):
            quote_text = line_stripped.replace(">", "").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(quote_text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(110, 110, 110) # 밝은 회색 처리
            i += 1
            continue

        # 7. 일반 문단 및 불릿 기호 처리
        # 마크다운 불릿(-)을 워드 리스트 스타일로 연계
        if line_stripped.startswith("- "):
            bullet_text = line_stripped.replace("- ", "", 1).strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(bullet_text)
            run.font.name = 'Malgun Gothic'
        elif line_stripped.startswith("* "):
            bullet_text = line_stripped.replace("* ", "", 1).strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(bullet_text)
            run.font.name = 'Malgun Gothic'
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line_stripped)
            run.font.name = 'Malgun Gothic'

        i += 1

    # 최종 문서 저장
    doc.save(docx_path)
    print(f"성공: Word 문서 파일이 성공적으로 저장되었습니다: {docx_path}")

if __name__ == "__main__":
    convert_md_to_docx()
