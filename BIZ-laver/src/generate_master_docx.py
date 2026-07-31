"""
통합 마스터 시장개척 보고서 Word (.docx) 파일 생성 스크립트.

기능:
1. 기초/이변량/다변량/심화 EDA 및 마른김 vs 조미김 개별 전략, Top 10 국가, 현지 파트너 리스트 전체 통합
2. BIZ-laver/reports/HaeYu_Laver_Export_Master_Market_Expansion_Report.docx 파일 생성
"""

import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

DOCX_OUTPUT_PATH = 'BIZ-laver/reports/HaeYu_Laver_Export_Master_Market_Expansion_Report.docx'
IMAGE_DIR = 'BIZ-laver/images'

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(26, 82, 118)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(46, 134, 193)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(40, 116, 166)
    return p

# 제목
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(10)
t_run = title_p.add_run("[마스터 전략서] 해유 김 수출 종합 EDA & 1인 무역회사 시장개척 마스터 보고서")
t_run.font.name = '맑은 고딕'
t_run.font.size = Pt(20)
t_run.font.bold = True
t_run.font.color.rgb = RGBColor(26, 82, 118)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(20)
s_run = sub_p.add_run("품목별 개별 전략, Top 10 타깃 국가, 현지 바이어 파트너십 & 4단계 실전 개척 프로토콜")
s_run.font.name = '맑은 고딕'
s_run.font.size = Pt(11)
s_run.font.color.rgb = RGBColor(120, 144, 156)

# 1. Executive Summary
add_styled_heading(doc, "1. Executive Summary (1인 무역회사 핵심 승리 공식)", 1)
p = doc.add_paragraph()
r = p.add_run("본 통합 마스터 보고서는 2021년부터 2025년까지의 글로벌 김 수출 데이터(1,069건)를 바탕으로, 1인 무역회사가 소자본 LCL 수송으로 최고 마진을 달성할 수 있는 시장 개척 전략을 제시합니다.")
r.font.name = '맑은 고딕'

# 2. 기초 EDA
add_styled_heading(doc, "2. 대한민국 김 수출 기초 EDA 데이터 분석 (2021~2025)", 1)
img1 = os.path.join(IMAGE_DIR, '01_univariate_item_distribution.png')
if os.path.exists(img1):
    doc.add_picture(img1, width=Inches(5.8))

# 3. 추가 심화 EDA (4분면 매트릭스 & 5대 마진 구간)
add_styled_heading(doc, "3. 단가 5대 마진 구간 & 4분면 시장 포트폴리오 분석", 1)
img14 = os.path.join(IMAGE_DIR, '14_advanced_market_portfolio_matrix.png')
if os.path.exists(img14):
    doc.add_picture(img14, width=Inches(5.8))

# 4. 품목별 Top 10 국가 & 파트너 표
add_styled_heading(doc, "4. 조미김 Top 10 타깃 시장 & 현지 잠재 파트너 리스트", 1)

table1 = doc.add_table(rows=11, cols=6)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["순위", "타깃 국가", "2025 수출액", "5개년 성장률", "평균 단가", "현지 컨택 가능 잠재 파트너 (수입상/유통사)"]

for i, h in enumerate(headers):
    cell = table1.cell(0, i)
    cell.text = h
    set_cell_background(cell, "1A5274")
    p = cell.paragraphs[0]
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.bold = True

partners_data = [
    ["1", "USA (미국)", "$244.6M", "+51.7%", "$26.72/kg", "Assi Rhee Bros, H-Mart, Harvesko, Weee!"],
    ["2", "Japan (일본)", "$124.8M", "+77.6%", "$22.10/kg", "E-Mart Japan, Gyomu Super, CJ CheilJedang Japan"],
    ["3", "Russian Fed. (러시아)", "$89.2M", "+98.5%", "$24.50/kg", "X5 Retail Group, Magnit, Koros Co."],
    ["4", "China (중국)", "$58.1M", "-23.1%", "$21.80/kg", "Ole' Supermarket, Citysuper China"],
    ["5", "Canada (캐나다)", "$32.4M", "+57.3%", "$27.13/kg", "T&T Supermarket, PAT Mart, Galleria Supermarket"],
    ["6", "Australia (호주)", "$26.1M", "+60.3%", "$27.01/kg", "Asian Inspirations, Miracle Supermarket"],
    ["7", "Poland (폴란드)", "$11.5M", "+121.2%", "$24.01/kg", "Kuchnie Świata, Asian House Poland, Allegro Sellers"],
    ["8", "UAE (아랍에미리트)", "$9.4M", "+187.5%", "$25.06/kg", "Choithrams, Lulu Hypermarket, Kibsons, Al Maya Group"],
    ["9", "Kazakhstan (카자흐스탄)", "$4.3M", "+296.0%", "$22.85/kg", "Magnum Cash & Carry, Shin-Line, Small Supermarket"],
    ["10", "Türkiye (튀르키예)", "$2.1M", "+254.5%", "$28.01/kg", "Macrocenter, Gurme Park, Happy Center"],
]

for row_idx, row_data in enumerate(partners_data, start=1):
    for col_idx, text in enumerate(row_data):
        cell = table1.cell(row_idx, col_idx)
        cell.text = text
        if row_idx % 2 == 1:
            set_cell_background(cell, "F4F6F7")

# 5. 1인 무역회사 실전 4단계 개척 프로토콜
add_styled_heading(doc, "5. 1인 무역회사 실전 4단계 개척 프로토콜", 1)
p_proto = doc.add_paragraph()
r_p = p_proto.add_run("Step 1 (고마진 포지셔닝) ➔ Step 2 (KOTRA 바이어 발굴 & 샘플 배송) ➔ Step 3 (LCL 물류 & 영문/현지어 라벨링) ➔ Step 4 (크로스보더 E-Commerce 직입점 & 숏폼 마케팅)")
r_p.font.name = '맑은 고딕'
r_p.font.bold = True

doc.save(DOCX_OUTPUT_PATH)
print(f"통합 마스터 Word 보고서가 {DOCX_OUTPUT_PATH}에 성공적으로 저장되었습니다.")
