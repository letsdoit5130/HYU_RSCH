"""
마크다운 마스터 보고서(HaeYu_Laver_Export_Master_Market_Expansion_Report.md)를
정교한 Word (.docx) 표준 양식 문서로 변환 및 생성하는 스크립트 (v2).
"""

import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

DOCX_OUTPUT_PATH = 'BIZ-laver/reports/HaeYu_Laver_Export_Master_Market_Expansion_Report_v2.docx'
IMAGE_DIR = 'BIZ-laver/images'

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(26, 82, 118)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 134, 193)
    return p

def add_body_p(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(10)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(44, 62, 80)
    return p

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(10)
title_p.paragraph_format.space_after = Pt(6)
t_run = title_p.add_run("[마스터 전략서] 해유 김 수출 종합 EDA & 1인 무역회사 글로벌 시장 개척 보고서 (v2)")
t_run.font.name = '맑은 고딕'
t_run.font.size = Pt(20)
t_run.font.bold = True
t_run.font.color.rgb = RGBColor(26, 82, 118)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(18)
s_run = sub_p.add_run("대한민국 김 수출 데이터(2021~2025) 기반 4분면 시장 포트폴리오, 단가 마진 구조 & 4단계 개척 프로토콜")
s_run.font.name = '맑은 고딕'
s_run.font.size = Pt(11)
s_run.font.color.rgb = RGBColor(120, 144, 156)

add_body_p(doc, "• 작성일자: 2026년 7월 24일", bold=True)
add_body_p(doc, "• 분석 대상: 대한민국 김 수출 통계 데이터 1,069건 (HS 121221 마른김 / HS 200899 조미김)", bold=True)
add_body_p(doc, "• 작성 목적: 1인 무역회사 사장님의 소자본 LCL 수송 및 고마진 신규 유망 시장 실전 개척", bold=True)

add_heading_1(doc, "1. Executive Summary (1인 무역회사 핵심 승리 공식)")
add_body_p(doc, "1. 조미김 (HS 200899) 100% 사격 집중: 2025년 조미김 평균 수출 단가는 $30.41/kg으로 마른김($21.56/kg) 대비 +41.1% 고마진을 형성하고 있습니다. 전체 수출액의 69.3%가 $20~$30/kg 구간, 9.3%가 $30~$50/kg 구간에 위치하여 1인 무역회사에 최적화되어 있습니다.")
add_body_p(doc, "2. Star Market 4개국 집중 공략: UAE(+187.5% 성장, $25.06/kg), 폴란드(+121.2% 성장, $24.01/kg), 콜롬비아(+156.7% 성장, $27.47/kg), 튀르키예(+254.5% 성장, $28.01/kg)를 핵심 타깃으로 설정합니다.")
add_body_p(doc, "3. 단가 변동성 리스크 대응: 고변동성 국가(CV > 40%, UAE/튀르키예/미국)는 FOB 고정단가 계약을, 저변동성 국가(CV < 25%, 폴란드/카자흐스탄)는 분기별 LCL 연속 출하를 적용합니다.")

add_heading_1(doc, "2. 대한민국 김 수출 기초 EDA 데이터 분석 (2021~2025)")

add_heading_2(doc, "2.1 수치형 변수 요약 통계량 (1,069 레코드)")
table_stat = doc.add_table(rows=4, cols=8)
table_stat.alignment = WD_TABLE_ALIGNMENT.CENTER
stat_headers = ["변수명", "평균 (Mean)", "표준편차", "최소값", "25% (Q1)", "중앙값", "75% (Q3)", "최대값"]
for i, h in enumerate(stat_headers):
    cell = table_stat.cell(0, i)
    cell.text = h
    set_cell_background(cell, "1A5274")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    cell.paragraphs[0].runs[0].font.bold = True

stat_data = [
    ["수출 금액 ($)", "$4,510,757", "$17,046,310", "$1.00", "$47,698", "$355,571", "$2,298,908", "$274,625,972"],
    ["수출 물량 (t)", "240.90 t", "733.91 t", "0.001 t", "2.10 t", "18.20 t", "131.70 t", "7,654.50 t"],
    ["단가 ($/kg)", "$22.95", "$17.56", "$0.03", "$13.57", "$21.57", "$29.74", "$242.75"]
]
for r_idx, row in enumerate(stat_data, start=1):
    for c_idx, val in enumerate(row):
        c = table_stat.cell(r_idx, c_idx)
        c.text = val
        if r_idx % 2 == 1:
            set_cell_background(c, "F4F6F7")

add_heading_2(doc, "2.2 품목별 (마른김 vs 조미김) 기본 비교")
img1 = os.path.join(IMAGE_DIR, '01_univariate_item_distribution.png')
if os.path.exists(img1):
    doc.add_picture(img1, width=Inches(5.8))

add_heading_1(doc, "3. 단가 5대 마진 구간 & 4분면 시장 포트폴리오 분석")

img13 = os.path.join(IMAGE_DIR, '13_advanced_price_bracket_distribution.png')
if os.path.exists(img13):
    doc.add_paragraph("그림 1: 품목별 수출 단가 5대 마진 구간 비중 (%)")
    doc.add_picture(img13, width=Inches(5.8))

img14 = os.path.join(IMAGE_DIR, '14_advanced_market_portfolio_matrix.png')
if os.path.exists(img14):
    doc.add_paragraph("\n그림 2: 글로벌 40개국 4분면 시장 포트폴리오 매트릭스")
    doc.add_picture(img14, width=Inches(5.8))

add_heading_1(doc, "4. 조미김 (HS 200899) Top 10 타깃 시장 & 현지 잠재 파트너")

table_partner = doc.add_table(rows=11, cols=6)
table_partner.alignment = WD_TABLE_ALIGNMENT.CENTER
p_headers = ["순위", "타깃 국가", "2025 수출액", "5개년 성장률", "평균 단가", "현지 컨택 가능 잠재 파트너"]
for i, h in enumerate(p_headers):
    cell = table_partner.cell(0, i)
    cell.text = h
    set_cell_background(cell, "1A5274")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    cell.paragraphs[0].runs[0].font.bold = True

partners_data = [
    ["1", "USA (미국)", "$244.6M", "+51.7%", "$26.72/kg", "Assi Rhee Bros, H-Mart, Harvesko, Weee!"],
    ["2", "Japan (일본)", "$124.8M", "+77.6%", "$22.10/kg", "E-Mart Japan, Gyomu Super, CJ CheilJedang Japan"],
    ["3", "Russian Fed. (러시아)", "$89.2M", "+98.5%", "$24.50/kg", "X5 Retail Group, Magnit, Koros Co."],
    ["4", "China (중국)", "$58.1M", "-23.1%", "$21.80/kg", "Ole' Supermarket, Citysuper China"],
    ["5", "Canada (캐나다)", "$32.4M", "+57.3%", "$27.13/kg", "T&T Supermarket, PAT Mart, Galleria Supermarket"],
    ["6", "Australia (호주)", "$26.1M", "+60.3%", "$27.01/kg", "Asian Inspirations, Miracle Supermarket"],
    ["7", "Poland (폴란드)", "$11.5M", "+121.2%", "$24.01/kg", "Kuchnie Świata, Asian House Poland, Allegro Sellers"],
    ["8", "UAE (아랍에미리트)", "$9.4M", "+187.5%", "$25.06/kg", "Choithrams, Lulu Hypermarket, Kibsons, Al Maya Group"],
    ["9", "Kazakhstan (카자흐스탄)", "$4.3M", "+296.0%", "$22.85/kg", "Magnum Cash & Carry, Shin-Line, Small Supermarket"],
    ["10", "Türkiye (튀르키예)", "$2.1M", "+254.5%", "$28.01/kg", "Macrocenter, Gurme Park, Happy Center"],
]

for r_idx, row in enumerate(partners_data, start=1):
    for c_idx, val in enumerate(row):
        c = table_partner.cell(r_idx, c_idx)
        c.text = val
        if r_idx % 2 == 1:
            set_cell_background(c, "F4F6F7")

add_heading_1(doc, "5. 1인 무역회사 실전 4단계 개척 프로토콜")
add_body_p(doc, "Step 1: 고마진 품목 포지셔닝 - 조미김 완제품(HS 200899) 스낵 파우치 + 비건/할랄 인증 + 특수 시즈닝", bold=True)
add_body_p(doc, "Step 2: KOTRA 및 B2B 바이어 발굴 - 두바이/알마티 무역관 지사화 사업 + 무료 샘플 항공 배송", bold=True)
add_body_p(doc, "Step 3: 소량 LCL 물류 & 영문/현지어 라벨링 - 초기 1~2 펠릿 LCL 수송으로 재고 최소화", bold=True)
add_body_p(doc, "Step 4: 크로스보더 E-Commerce & 숏폼 마케팅 - Amazon, Noon, Allegro 입점 + 틱톡 K-Food 숏폼", bold=True)

doc.save(DOCX_OUTPUT_PATH)
print(f"Master Word v2 파일이 {DOCX_OUTPUT_PATH}에 정상 저장되었습니다.")
