"""
kyobooks_harness Word 비즈니스 보고서 생성 스크립트

이 모듈은 python-docx 라이브러리를 사용하여 수집 및 분석된 
데이터를 바탕으로 고품질 오피스 Word 보고서(.docx)를 생성합니다.

작성일: 2026-07-19
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_docx_report():
    print("[Word Report] Word 비즈니스 보고서 생성을 시작합니다...")
    
    current_dir = os.path.dirname(os.path.abspath(__file__))
    data_path = os.path.abspath(os.path.join(current_dir, "..", "data", "raw_data.csv"))
    output_path = os.path.abspath(os.path.join(current_dir, "..", "docs", "report.docx"))
    
    if not os.path.exists(data_path):
        print(f"[Word Report] [ERROR] 원시 데이터가 존재하지 않습니다: {data_path}")
        return
        
    df = pd.read_csv(data_path)
    
    # Document 개체 생성
    doc = Document()
    
    # 1. 문서 제목
    title = doc.add_paragraph()
    title_run = title.add_run("kyobooks_harness 종합 비즈니스 보고서")
    title_run.font.name = '맑은 고딕'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. 문서 개요 및 정보
    doc.add_paragraph("본 보고서는 범용 크롤링 및 EDA 분석 스킬을 활용하여 자동으로 빌드된 분석 리포트입니다.")
    
    # 3. 섹션 1: 데이터 기초 통계
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. 수집 데이터 개요")
    h1_run.font.name = '맑은 고딕'
    h1_run.font.bold = True
    
    doc.add_paragraph(f"수집된 원시 데이터는 총 {len(df)}건이며, 주요 필드는 상품번호, 순위, 도서명, 저자, 출판사, 출판일, 정가, 할인가, 할인율, 평점 등 입니다.")
    
    # 데이터 테이블 추가
    table = doc.add_table(rows=1, cols=min(len(df.columns), 5))
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    
    # 헤더 작성
    for idx, col_name in enumerate(df.columns[:5]):
        hdr_cells[idx].text = str(col_name)
        
    # 데이터 일부(상위 5개) 채우기
    for r_idx in range(min(len(df), 5)):
        row_cells = table.add_row().cells
        for c_idx, col_name in enumerate(df.columns[:5]):
            row_cells[c_idx].text = str(df.iloc[r_idx, c_idx])
            
    doc.add_paragraph() # 문단 띄우기
    
    # 4. 섹션 2: 시각화 이미지 포함
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. 시각화 분석 결과")
    h2_run.font.name = '맑은 고딕'
    h2_run.font.bold = True
    
    chart_path = os.path.abspath(os.path.join(current_dir, "..", "images", "summary_chart.png"))
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(5.5))
        doc.add_paragraph("그림 1. 주요 수집 항목 요약 차트")
    else:
        doc.add_paragraph("[차트 이미지를 찾을 수 없습니다]")
        
    doc.save(output_path)
    print(f"[Word Report] Word 보고서 저장 완료: {output_path}")
    return output_path

if __name__ == "__main__":
    build_docx_report()
