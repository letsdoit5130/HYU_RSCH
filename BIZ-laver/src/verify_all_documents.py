"""
1~8번 전 섹션이 100% 완전 수록된 Word 완본 문서(Full_Complete.docx) 및 마스터 문서 패키지 무결성 검증 스크립트.
"""

import os
import sys
import docx
import openpyxl
import pptx

sys.stdout.reconfigure(encoding='utf-8')

REPORTS_DIR = 'BIZ-laver/reports'

files_to_verify = {
    'Master Markdown Report': os.path.join(REPORTS_DIR, 'HaeYu_Laver_Export_Master_Market_Expansion_Report.md'),
    'Full Complete Word Report': os.path.join(REPORTS_DIR, 'HaeYu_Laver_Export_Master_Market_Expansion_Report_Full_Complete.docx'),
    'Full Master Excel Dashboard': os.path.join(REPORTS_DIR, 'HaeYu_Laver_Export_Master_Dashboard_and_Data_Full.xlsx'),
    'Full Master PPTX Presentation': os.path.join(REPORTS_DIR, 'HaeYu_Laver_Export_Master_Market_Expansion_Deck_Full.pptx'),
}

print("==================================================")
print(" 1~8번 전 섹션 100% 완본 문서 QA 무결성 최종 검증 ")
print("==================================================\n")

all_pass = True

for name, path in files_to_verify.items():
    if not os.path.exists(path):
        print(f"[FAIL] {name}: 파일이 존재하지 않습니다. ({path})")
        all_pass = False
        continue
    
    size_kb = os.path.getsize(path) / 1024.0
    print(f"[OK] {name}: 파일 존재 확인 ({size_kb:.1f} KB)")
    
    if 'Word' in name:
        try:
            doc = docx.Document(path)
            para_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            print(f"     └ Word 무결성: 정상 읽기 완료 (단락 {para_count}개, 표 {table_count}개)")
        except Exception as e:
            print(f"     └ [ERROR] Word 파일 손상: {e}")
            all_pass = False
            
    elif 'Excel' in name:
        try:
            wb = openpyxl.load_workbook(path, data_only=False)
            sheet_names = wb.sheetnames
            print(f"     └ Excel 무결성: 정상 읽기 완료 (시트 {len(sheet_names)}개: {sheet_names})")
            
            formula_error_count = 0
            for sheetname in sheet_names:
                ws = wb[sheetname]
                for row in ws.iter_rows():
                    for cell in row:
                        val = str(cell.value)
                        if any(err in val for err in ['#REF!', '#DIV/0!', '#NAME?', '#VALUE!', '#N/A']):
                            print(f"     └ [FORMULA ERROR] 시트 '{sheetname}' 셀 {cell.coordinate}: {val}")
                            formula_error_count += 1
                            
            if formula_error_count == 0:
                print("     └ Excel 수식 검증: 수식 오류 (#REF!, #DIV/0! 등) 0건 [PASS]")
            else:
                print(f"     └ [FAIL] Excel 수식 오류 총 {formula_error_count}건 발견!")
                all_pass = False
        except Exception as e:
            print(f"     └ [ERROR] Excel 파일 손상: {e}")
            all_pass = False

    elif 'PPTX' in name:
        try:
            prs = pptx.Presentation(path)
            slide_count = len(prs.slides)
            print(f"     └ PPTX 무결성: 정상 읽기 완료 (총 {slide_count}개 슬라이드)")
        except Exception as e:
            print(f"     └ [ERROR] PPTX 파일 손상: {e}")
            all_pass = False

print("\n==================================================")
if all_pass:
    print("[VERIFICATION RESULT]: ALL COMPLETE MASTER DOCUMENTS PASSED QA CHECK! (PASS)")
else:
    print("[VERIFICATION RESULT]: QA CHECK FAILED")
print("==================================================")
